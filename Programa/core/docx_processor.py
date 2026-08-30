"""core/docx_processor.py
Extracción y clasificación de bloques a partir de un documento Word (.docx).

A diferencia del PDF (que es un producto "aplanado" y obliga a adivinar la
estructura con heurísticas tipográficas), el .docx conserva estructura
semántica explícita: **estilos de párrafo** (Título, Título 1, Epígrafe…),
cursiva/negrita reales por run y tablas/figuras como objetos. Por eso aquí la
clasificación es **guiada por estilos** primero, y solo cae en reglas de
contenido como respaldo.

Devuelve la MISMA forma que core.pdf_processor.procesar_pdf, para que el resto
del programa (endpoints, exportadores) lo consuma sin cambios:

    resultado = procesar_docx("articulo.docx")
    resultado["bloques"]                # list[dict] {contenido, clasificacion, size, bold, italic}
    resultado["figuras"]                # list[dict] {ruta, pie, ancla, pagina, origen}
    resultado["tablas"]                 # list[dict] {ruta, hoja, rotulo, descripcion, ancla, pagina, origen}
    resultado["body_size"]              # int  (tamaño de fuente dominante, en pt)
    resultado["resumen"]                # str
    resultado["fig_dir"]                # str | None  (temporal de imágenes; limpiar con rmtree)
    resultado["metadatos_detectados"]   # dict
"""

from __future__ import annotations

import os
import re
import tempfile
from collections import Counter
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from core.constans import CLASE_COMPAT as _CLASE_COMPAT
from core.utils import (
    es_como_citar as _es_como_citar,
    es_fecha_mss as _es_fecha_mss,
    es_doi as _es_doi,
    limpiar_prefijo_pie_figura as _limpiar_prefijo_pie_figura,
    extraer_fechas_mss as _extraer_fechas_mss,
    fecha_mss_a_iso as _fecha_mss_a_iso,
    extraer_doi as _extraer_doi,
    extraer_volumen_pagina as _extraer_volumen_pagina,
    extraer_issn as _extraer_issn,
    es_linea_masthead as _es_linea_masthead,
    es_linea_filiacion as _es_linea_filiacion,
)
from core.zonas import detectar_breakpoints, CursorZonas, Breakpoint

# ─────────────────────────────────────────────────────────────────────────────
# Clasificación
# ─────────────────────────────────────────────────────────────────────────────

_SECCIONES_EXACTAS = {
    "resumen", "abstract", "resumen no técnico", "non-technical abstract",
    "palabras clave", "keywords", "introducción", "introduction",
    "conclusiones", "conclusions", "referencias", "references",
    "agradecimientos", "acknowledgements", "acknowledgments",
    "discusión", "discussion", "metodología", "methods", "resultados", "results",
    "contribuciones de los autores", "contribucción de autores",
    "contribución de autores", "authors' contribution",
    "conflicto de intereses", "conflict of interest",
    "conflicts of interest", "competing interests",
    "paleontología sistemática", "systematic palaeontology",
}

_RE_PALABRAS_CLAVE = re.compile(r"^(palabras\s+clave|keywords)\s*[:\.]", re.IGNORECASE)
_RE_TITULO_TABLA   = re.compile(r"^(tabla|table)\s+\d+[\.\:\s]", re.IGNORECASE)
_RE_PIE_FIGURA     = re.compile(r"^(figura|fig\.?|figure)\s+\d+[a-z]?[\.\:\s]", re.IGNORECASE)
_RE_NIVEL1         = re.compile(r"^\d+\.\s+\S")          # "1. Introducción"
_RE_NIVEL2         = re.compile(r"^\d+\.\d+")            # "1.2 Subsección"


def _clase_por_estilo(estilo: str) -> str | None:
    """Mapea un nombre de estilo de Word a la clasificación semántica del editor.

    Cubre nombres en inglés (Word en-US) y español (Word es-ES). Ojo con el
    español: «Título» (sin número) = Title, «Título 1» = Heading 1.
    """
    sn = (estilo or "").strip().lower()
    if not sn:
        return None
    ultimo = sn[-1]
    tiene_num = ultimo.isdigit()

    # Título del documento
    if sn in ("title", "título", "titulo") or sn.startswith("title") and not tiene_num:
        return "Título principal"
    if sn in ("subtitle", "subtítulo", "subtitulo"):
        return "Título secundario"

    # Encabezados de sección (Heading N / Título N)
    if sn.startswith(("heading", "título", "titulo")) and tiene_num:
        return "Subencabezado" if ultimo == "1" else "Subencabezado-bajo"

    # Epígrafes / leyendas de figura o tabla
    if sn.startswith(("caption", "epígrafe", "epigrafe", "leyenda")):
        return "Pie de figura"   # el texto decidirá si es tabla

    return None   # Normal / Body Text / Cita… → decide el contenido


def _clasificar(texto: str, estilo: str, bold: bool, italic: bool,
                body_size: int, size: int) -> str:
    """Clasifica un párrafo: estilo de Word primero, reglas de contenido después."""
    t = texto.strip()
    t_low = t.lower()

    # 1) Reglas de contenido fuertes (mandan sobre el estilo)
    if t_low in _SECCIONES_EXACTAS:
        return "Encabezado sección"
    if _es_linea_masthead(t):
        return "Encabezado sección"
    if _es_como_citar(t):
        return "Cómo citar"
    if _es_fecha_mss(t) or _es_doi(t):
        return "Fecha manuscrito"
    if _RE_PALABRAS_CLAVE.match(t):
        return "Palabras clave"
    if _RE_TITULO_TABLA.match(t):
        return "Título tabla"
    if _RE_PIE_FIGURA.match(t):
        return "Pie de figura"
    if _es_linea_filiacion(t):
        return "Filiación"

    # 2) Estilo de Word
    por_estilo = _clase_por_estilo(estilo)
    if por_estilo == "Pie de figura":
        # Un epígrafe cuyo texto empieza por «Tabla N» es título de tabla.
        return "Título tabla" if _RE_TITULO_TABLA.match(t) else "Pie de figura"
    if por_estilo:
        return por_estilo

    # 3) Respaldo por contenido para párrafos «Normal»
    if _RE_NIVEL1.match(t) and not _RE_NIVEL2.match(t):
        return "Subencabezado"
    if _RE_NIVEL2.match(t):
        return "Subencabezado-bajo"
    if re.search(r"@[\w\-\.]+\.\w{2,}", t) and len(t) < 120:
        return "Email / Metadatos"
    if len(t) < 3:
        return "Ignorar"
    return "Cuerpo"


def _clasificar_con_zonas(
    par_idx: int, texto: str, estilo: str, bold: bool, italic: bool,
    body_size: int, size: int, zona: Breakpoint | None,
) -> str:
    """Igual que _clasificar(), pero primero revisa en qué ZONA del artículo
    cae el párrafo (Resumen / Palabras clave / Referencias), usando el
    breakpoint más reciente detectado por core.zonas (que reconoce TODAS
    las ocurrencias de cada ancla, no solo la primera — ver core/zonas.py).
    Solo si el párrafo no cae en ninguna zona especial (zona is None, o cae
    en la zona genérica "cuerpo"), se recurre a la clasificación normal por
    estilo/contenido (_clasificar), que ya trae sus propios detectores de
    encabezados numerados, Cómo citar, fechas de manuscrito, etc.
    """
    if zona is None:
        return _clasificar(texto, estilo, bold, italic, body_size, size)

    es_ancla = zona.idx == par_idx

    if zona.tipo == "resumen":
        if es_ancla:
            return "Encabezado sección"
        if _es_como_citar(texto):
            return "Cómo citar"
        if _es_fecha_mss(texto) or _es_doi(texto):
            return "Fecha manuscrito"
        return "Cuerpo del abstract"

    if zona.tipo == "palabras_clave":
        if es_ancla and not zona.inline:
            return "Encabezado sección"
        if not es_ancla and _es_como_citar(texto):
            return "Cómo citar"
        if not es_ancla and (_es_fecha_mss(texto) or _es_doi(texto)):
            return "Fecha manuscrito"
        return "Palabras clave"

    if zona.tipo == "referencias":
        if es_ancla:
            return "Encabezado sección"
        if re.match(r"^(cómo citar|how to cite)", texto.strip().lower()):
            return "Cómo citar"
        if _es_fecha_mss(texto) or _es_doi(texto):
            return "Fecha manuscrito"
        return "Referencia"

    # zona.tipo == "cuerpo": el estilo de Word ya distingue encabezados de
    # párrafos normales de forma fiable, así que acá basta con la
    # clasificación normal (que también reconoce "N. Título" numerado).
    return _clasificar(texto, estilo, bold, italic, body_size, size)


# ─────────────────────────────────────────────────────────────────────────────
# Formato de párrafo (negrita / cursiva / tamaño dominante por caracteres)
# ─────────────────────────────────────────────────────────────────────────────

def _formato_parrafo(par: Paragraph) -> tuple[bool, bool, int | None]:
    runs = [r for r in par.runs if r.text and r.text.strip()]
    total = sum(len(r.text) for r in runs) or 1
    bold_chars = sum(len(r.text) for r in runs if r.bold)
    ital_chars = sum(len(r.text) for r in runs if r.italic)
    sizes: list[float] = []
    for r in runs:
        try:
            if r.font is not None and r.font.size is not None:
                sizes.append(r.font.size.pt)
        except Exception:
            pass
    size = round(sum(sizes) / len(sizes)) if sizes else None
    return (bold_chars / total >= 0.6, ital_chars / total >= 0.6, size)


# ─────────────────────────────────────────────────────────────────────────────
# Imágenes embebidas
# ─────────────────────────────────────────────────────────────────────────────

_CT_EXT = {
    "image/png": "png", "image/jpeg": "jpg", "image/jpg": "jpg",
    "image/gif": "gif", "image/bmp": "bmp", "image/tiff": "tif",
    "image/x-emf": "emf", "image/x-wmf": "wmf", "image/webp": "webp",
}


def _imagenes_de_parrafo(par: Paragraph, doc, fig_dir: str,
                         contador: list[int]) -> list[dict]:
    """Extrae las imágenes inline de un párrafo a fig_dir. Devuelve figuras
    (con pie vacío; el epígrafe cercano lo rellena después)."""
    figuras: list[dict] = []
    blips = par._p.findall(".//" + qn("a:blip"))
    for blip in blips:
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not rid:
            continue
        try:
            parte = doc.part.related_parts[rid]
            blob = parte.blob
        except Exception:
            continue
        ext = _CT_EXT.get(getattr(parte, "content_type", ""), "")
        if not ext:
            ext = (os.path.splitext(getattr(parte, "partname", "") or "")[1]
                   .lstrip(".").lower() or "png")
        contador[0] += 1
        nom = f"img_{contador[0]:03d}.{ext}"
        ruta = os.path.join(fig_dir, nom)
        try:
            with open(ruta, "wb") as f:
                f.write(blob)
        except Exception:
            continue
        figuras.append({
            "ruta": ruta, "pie": "", "ancla": "",
            "pagina": contador[0], "origen": "auto_docx",
        })
    return figuras


# ─────────────────────────────────────────────────────────────────────────────
# Tablas
# ─────────────────────────────────────────────────────────────────────────────

def _tabla_a_matriz(tabla: Table) -> list[list[str]]:
    filas: list[list[str]] = []
    for fila in tabla.rows:
        filas.append([(c.text or "").strip() for c in fila.cells])
    return filas


def _guardar_tabla_xlsx(matriz: list[list[str]], tab_dir: str,
                        idx: int) -> dict | None:
    if not matriz:
        return None
    try:
        import openpyxl as xl
    except ImportError:
        return None
    hoja = f"Tabla_{idx}"
    ruta = os.path.join(tab_dir, f"tabla_docx_{idx:02d}.xlsx")
    try:
        wb = xl.Workbook()
        ws = wb.active
        ws.title = hoja
        for fila in matriz:
            ws.append([c if c is not None else "" for c in fila])
        wb.save(ruta)
        wb.close()
    except Exception:
        return None
    return {
        "ruta": ruta, "hoja": hoja,
        "rotulo": "", "descripcion": "", "ancla": "",
        "pagina": idx, "origen": "auto_docx",
    }


# ─────────────────────────────────────────────────────────────────────────────
# Metadatos del artículo (volumen, número, DOI, fechas…)
# ─────────────────────────────────────────────────────────────────────────────

def _extraer_metadatos(bloques: list[dict]) -> dict[str, Any]:
    m: dict[str, Any] = {
        "volumen": "", "numero": "", "anio": "",
        "pagina_inicio": "", "pagina_fin": "", "doi": "", "issn": "",
        "fecha_recibido": "", "fecha_corregido": "", "fecha_aceptado": "",
        "fecha_recibido_iso": "", "fecha_corregido_iso": "", "fecha_aceptado_iso": "",
    }
    for b in bloques:
        texto = b.get("contenido", "")
        if not texto:
            continue
        if not m["volumen"]:
            vol = _extraer_volumen_pagina(texto)
            if vol:
                m.update({k: v for k, v in vol.items() if v})
        if not m["issn"]:
            issn = _extraer_issn(texto)
            if issn:
                m["issn"] = issn
        if b.get("clasificacion") == "Fecha manuscrito" or _es_fecha_mss(texto) or _es_doi(texto):
            if not m["doi"]:
                doi = _extraer_doi(texto)
                if doi:
                    m["doi"] = doi
            for clave, valor in _extraer_fechas_mss(texto).items():
                campo = f"fecha_{clave}"
                if not m.get(campo):
                    m[campo] = valor
                    iso = _fecha_mss_a_iso(valor)
                    if iso:
                        m[f"{campo}_iso"] = iso
    return m


# ─────────────────────────────────────────────────────────────────────────────
# Iteración del cuerpo en orden (párrafos y tablas intercalados)
# ─────────────────────────────────────────────────────────────────────────────

def _iter_cuerpo(doc):
    """Recorre el cuerpo del documento en orden real, devolviendo Paragraph y
    Table (python-docx los expone en listas separadas, perdiendo el orden)."""
    for hijo in doc.element.body.iterchildren():
        if hijo.tag == qn("w:p"):
            yield Paragraph(hijo, doc)
        elif hijo.tag == qn("w:tbl"):
            yield Table(hijo, doc)


# ─────────────────────────────────────────────────────────────────────────────
# Pipeline principal
# ─────────────────────────────────────────────────────────────────────────────

def procesar_docx(ruta: str) -> dict[str, Any]:
    doc = Document(ruta)

    base = re.sub(r"[^A-Za-z0-9_-]+", "_",
                  os.path.splitext(os.path.basename(ruta))[0]).strip("_")[:24] or "docx"
    try:
        fig_dir = tempfile.mkdtemp(prefix=f"pm_docx_fig_{base}_")
    except Exception:
        fig_dir = os.path.join(os.getcwd(), f"_pm_docx_fig_{base}")
        os.makedirs(fig_dir, exist_ok=True)
    try:
        tab_dir = tempfile.mkdtemp(prefix=f"pm_docx_tab_{base}_")
    except Exception:
        tab_dir = os.path.join(os.getcwd(), f"_pm_docx_tab_{base}")
        os.makedirs(tab_dir, exist_ok=True)

    # Tamaño de fuente dominante (para el respaldo de clasificación / info).
    sizes: list[int] = []
    for par in doc.paragraphs:
        for r in par.runs:
            try:
                if r.font is not None and r.font.size is not None:
                    sizes.append(round(r.font.size.pt))
            except Exception:
                pass
    body_size = Counter(sizes).most_common(1)[0][0] if sizes else 12

    # ── Pre-escaneo: detectar TODAS las anclas semánticas del artículo ────────
    # (Resumen, Palabras clave, encabezados del cuerpo, Referencias), con la
    # misma lógica que PDF (core/zonas.py) — reconoce cada ocurrencia, no
    # solo la primera, así que artículos con resumen+abstract+resumen no
    # técnico+non-technical abstract (cada uno con su propio "Palabras
    # clave"/"Keywords") se delimitan bien. doc.paragraphs conserva el mismo
    # orden que _iter_cuerpo recorre los párrafos (ambos son los w:p de nivel
    # superior del body), así que el índice par_idx usado más abajo es
    # comparable con estos.
    parrafos_texto = [(p.text or "").strip() for p in doc.paragraphs]
    breakpoints = detectar_breakpoints(parrafos_texto)
    cursor_zonas = CursorZonas(breakpoints)

    bloques: list[dict] = []
    figuras: list[dict] = []
    tablas:  list[dict] = []
    contador_img = [0]
    n_tabla = 0
    par_idx = -1

    for elem in _iter_cuerpo(doc):
        if isinstance(elem, Table):
            n_tabla += 1
            t = _guardar_tabla_xlsx(_tabla_a_matriz(elem), tab_dir, n_tabla)
            if t:
                tablas.append(t)
            continue

        par: Paragraph = elem
        par_idx += 1

        # Imágenes inline de este párrafo (antes de decidir si hay texto).
        figs = _imagenes_de_parrafo(par, doc, fig_dir, contador_img)
        figuras.extend(figs)

        texto = (par.text or "").strip()
        if not texto:
            continue

        estilo = par.style.name if par.style is not None else ""
        bold, italic, size = _formato_parrafo(par)
        zona = cursor_zonas.avanzar(par_idx)
        cls = _clasificar_con_zonas(
            par_idx, texto, estilo, bold, italic, body_size,
            size if size is not None else body_size,
            zona,
        )
        if cls == "Ignorar":
            continue

        # Un epígrafe de figura rellena el pie de la última figura sin pie.
        if cls == "Pie de figura":
            for fig in reversed(figuras):
                if not fig["pie"]:
                    fig["pie"] = _limpiar_prefijo_pie_figura(texto)
                    break

        bloques.append({
            "contenido": texto,
            "clasificacion": _CLASE_COMPAT.get(cls, cls),
            "size": size if size is not None else body_size,
            "bold": bold,
            "italic": italic,
        })

    metadatos = _extraer_metadatos(bloques)

    conteo = Counter(b["clasificacion"] for b in bloques)
    resumen = "  |  ".join(f"{k}: {v}" for k, v in conteo.most_common(6))
    resumen += f"  |  Figuras: {len(figuras)}  |  Tablas: {len(tablas)}"

    return {
        "bloques":              bloques,
        "figuras":              figuras,
        "tablas":                tablas,
        "body_size":             body_size,
        "resumen":               resumen,
        "fig_dir":               fig_dir,
        "tab_dir":               tab_dir,
        "metadatos_detectados": metadatos,
    }